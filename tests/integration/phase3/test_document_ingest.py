import hashlib
import uuid
from io import BytesIO

import pytest
from docx import Document
from sqlalchemy import delete, select

from energy_agent.observability.tracing import LocalTracer
from energy_agent.persistence.models import ManualChunkModel, ManualDocumentModel
from energy_agent.persistence.mysql import create_mysql_engine, create_session_factory
from energy_agent.persistence.repositories.manual_document import ManualDocumentRepository
from energy_agent.providers.milvus import MilvusVectorProvider
from energy_agent.providers.minio import MinioDocumentProvider
from energy_agent.retrieval.ingestion.manifests import (
    DocumentManifest,
    IndexStatus,
    ReviewStatus,
)
from energy_agent.retrieval.ingestion.service import DocumentIngestionService

pytestmark = pytest.mark.integration

MYSQL_DSN = "mysql+asyncmy://energy:energy_dev@localhost:3306/energy_agent"


class DeterministicTestEmbedding:
    model = "BAAI/bge-m3"
    dimension = 1024

    async def embed(self, texts: list[str]) -> list[list[float]]:
        vectors = []
        for text in texts:
            output = [0.0] * 1024
            output[int(hashlib.sha256(text.encode()).hexdigest()[:4], 16) % 1024] = 1.0
            vectors.append(output)
        return vectors


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("散热系统维护", level=1)
    document.add_paragraph("1. 检查散热风扇供电和转速。")
    document.add_paragraph("2. 检查滤网积尘与风道堵塞。")
    document.add_paragraph("注意：断电操作必须由授权人员执行。")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


@pytest.mark.asyncio
async def test_real_docx_ingest_minio_mysql_milvus_and_idempotency() -> None:
    suffix = uuid.uuid4().hex[:8]
    doc_id = f"DOC-PHASE3-{suffix}"
    engine = create_mysql_engine(MYSQL_DSN)
    factory = create_session_factory(engine)
    minio = MinioDocumentProvider(
        endpoint="localhost:9000",
        access_key="energy",
        secret_key="energy_minio_dev",
        bucket="energy-documents-test",
        secure=False,
    )
    milvus = MilvusVectorProvider(
        uri="http://localhost:19530",
        token=None,
        manual_collection=f"manual_ingest_{suffix}",
        ticket_collection=f"ticket_ingest_{suffix}",
        dimension=1024,
        metric_type="COSINE",
    )
    content = docx_bytes()
    manifest = DocumentManifest(
        doc_id=doc_id,
        document_name="pcs-maintenance.docx",
        version="1.0",
        content_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        device_type="PCS",
        device_model="SC5000",
        manufacturer="EnergyCo",
        alarm_name="温度告警",
        review_status=ReviewStatus.APPROVED,
        effective=True,
    )
    try:
        await minio.ensure_bucket()
        await milvus.ensure_collections()
        service = DocumentIngestionService(
            repository=ManualDocumentRepository(factory),
            minio=minio,
            milvus=milvus,
            embedding=DeterministicTestEmbedding(),  # type: ignore[arg-type]
            tracer=LocalTracer(),
            max_bytes=10_000_000,
        )
        result = await service.ingest(manifest, content)
        duplicate = await service.ingest(manifest, content)
        assert result.index_status == IndexStatus.INDEXED
        assert result.chunk_count >= 3
        assert duplicate.existing is True
        assert await minio.get(result.object_key) == content
        assert result.file_sha256 == hashlib.sha256(content).hexdigest()
        async with factory() as session:
            document = (
                await session.execute(
                    select(ManualDocumentModel).where(ManualDocumentModel.doc_id == doc_id)
                )
            ).scalar_one()
            chunks = (
                (
                    await session.execute(
                        select(ManualChunkModel).where(ManualChunkModel.doc_id == doc_id)
                    )
                )
                .scalars()
                .all()
            )
        assert document.index_status == "INDEXED"
        assert len(chunks) == result.chunk_count
        hits = await milvus.search(
            "manual",
            (await DeterministicTestEmbedding().embed([chunks[0].embedding_text or ""]))[0],
            [chunk.chunk_id for chunk in chunks],
            1,
        )
        assert hits
    finally:
        async with factory() as session, session.begin():
            await session.execute(delete(ManualChunkModel).where(ManualChunkModel.doc_id == doc_id))
            await session.execute(
                delete(ManualDocumentModel).where(ManualDocumentModel.doc_id == doc_id)
            )
        await milvus.close()
        await engine.dispose()
