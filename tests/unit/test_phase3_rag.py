from datetime import UTC, datetime, timedelta
from io import BytesIO

import pytest
from docx import Document
from pydantic import ValidationError
from pypdf import PdfWriter

from energy_agent.core.config import Settings
from energy_agent.core.errors import OcrRequiredError
from energy_agent.model.gateway import _openai_strict_schema
from energy_agent.retrieval.contracts import QueryRewrite, RetrievalCandidate, SourceType
from energy_agent.retrieval.dedup import deduplicate_and_diversify
from energy_agent.retrieval.evidence import build_evidence_package
from energy_agent.retrieval.ingestion.chunking import chunk_blocks
from energy_agent.retrieval.ingestion.index_tickets import build_ticket_embedding_text
from energy_agent.retrieval.ingestion.parsers import ParsedBlock, parse_document
from energy_agent.retrieval.keyword import LightweightKeywordRetriever
from energy_agent.retrieval.merge import merge_candidates
from energy_agent.retrieval.query_rewrite import rewrite_query, rule_rewrite
from energy_agent.retrieval.scoring import (
    final_score,
    freshness_score,
    retrieval_score,
    score_candidate,
    verification_score,
)
from energy_agent.retrieval.tokenization import tokenize
from energy_agent.tools.contracts import ManualSearchInput, TicketSearchInput, ToolContext


def candidate(
    identifier: str,
    *,
    source: SourceType = SourceType.MANUAL,
    score: float = 0.8,
    embedding: list[float] | None = None,
) -> RetrievalCandidate:
    return RetrievalCandidate(
        source_type=source,
        source_id=identifier,
        chunk_id=f"{identifier}-chunk" if source == SourceType.MANUAL else None,
        content_summary="PCS 温度告警，检查散热风扇和滤网",
        citation=f"[来源: {identifier}]",
        metadata={
            "verified": True,
            "closed": True,
            "effective": True,
            "alarm_name": "温度告警",
            "device_type": "PCS",
        },
        keyword_score=score,
        vector_score=score,
        rerank_score=score,
        final_score=score,
        embedding=embedding,
    )


def test_query_rewrite_normalizes_aliases_and_preserves_identifiers() -> None:
    result = rule_rewrite("SC5000 储能变流器过温，风机 A12 异常")
    assert result.device_type == "PCS"
    assert result.normalized_alarm_name == "温度告警"
    assert result.component == "散热风扇"
    assert {"散热", "散热风扇", "滤网", "风道"} <= set(result.keyword_terms)
    assert "SC5000" in result.keyword_terms
    assert result.rewrite_mode == "rules"


@pytest.mark.asyncio
async def test_model_rewrite_failure_falls_back_to_rules() -> None:
    async def broken(_: dict[str, object]) -> object:
        raise TimeoutError

    result = await rewrite_query("PCS 温度高", mode="model_enhanced", model_rewrite=broken)
    assert result.rewrite_mode == "rules"
    assert result.warnings == ["QUERY_REWRITE_FAILED"]


def test_tokenizer_and_bm25_reward_exact_alarm_and_title_matches() -> None:
    assert "pcs" in tokenize("PCS温度 Alarm-100")
    rows = [
        {
            "id": "strong",
            "alarm_name": "ALARM-100",
            "chapter_title": "PCS 散热维护",
            "content": "检查风扇",
        },
        {
            "id": "weak",
            "alarm_name": "通信告警",
            "chapter_title": "网络",
            "content": "PCS 正常",
        },
    ]
    ranked = LightweightKeywordRetriever().rank(
        "PCS ALARM-100 风扇",
        rows,
        ("alarm_name", "chapter_title", "content"),
        2,
        title_fields=("alarm_name", "chapter_title"),
        exact_terms=("ALARM-100", "PCS"),
    )
    assert ranked[0]["id"] == "strong"
    assert ranked[0]["keyword_raw_score"] > ranked[1]["keyword_raw_score"]


def test_txt_docx_and_scanned_pdf_parsing() -> None:
    assert parse_document("manual.txt", "中文 English".encode())[0].text == "中文 English"
    document = Document()
    document.add_heading("散热维护", level=1)
    document.add_paragraph("1. 检查风扇")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "告警"
    table.cell(0, 1).text = "动作"
    stream = BytesIO()
    document.save(stream)
    blocks = parse_document("manual.docx", stream.getvalue())
    assert any(block.section_type == "表格" for block in blocks)
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    pdf = BytesIO()
    writer.write(pdf)
    with pytest.raises(OcrRequiredError):
        parse_document("scan.pdf", pdf.getvalue())


def test_chunking_is_structural_overlapping_stable_and_lossless() -> None:
    text = "# 散热系统\n" + "温度持续升高时应检查风扇运行状态。" * 45
    blocks = [ParsedBlock(text=text, page_no=3)]
    first = chunk_blocks("DOC-1", "1.0", blocks, target=320, overlap=60)
    second = chunk_blocks("DOC-1", "1.0", blocks, target=320, overlap=60)
    assert len(first) > 1
    assert [item.chunk_id for item in first] == [item.chunk_id for item in second]
    assert all(item.chapter_title == "散热系统" for item in first)
    assert first[0].content[-40:] in first[1].content
    assert "温度持续升高" in "".join(item.content for item in first)
    assert chunk_blocks("EMPTY", "1", []) == []


def test_merge_scoring_freshness_verification_and_final_score() -> None:
    keyword = [candidate("DOC-1", score=0.7)]
    vector = [candidate("DOC-1", score=0.9), candidate("DOC-2", score=0.5)]
    merged = merge_candidates(keyword, vector)
    assert len(merged) == 2
    assert merged[0].keyword_score is not None and merged[0].vector_score is not None
    assert retrieval_score(1, 1, 1) == 1
    assert final_score(1, 1, 1, 1, 1) == 1
    now = datetime.now(UTC)
    assert freshness_score(SourceType.TICKET, now - timedelta(days=20), now=now) == 1
    assert freshness_score(SourceType.TICKET, now - timedelta(days=400), now=now) == 0.5
    assert verification_score(SourceType.TICKET, {"verified": False, "closed": False}) == 0.3
    scored = score_candidate(merged[0], {"alarm_name": "温度告警", "device_type": "PCS"})
    assert 0 < scored.final_score <= 1


def test_exact_semantic_dedup_limits_and_source_diversity() -> None:
    manual_a = candidate("DOC-1", score=0.95, embedding=[1, 0])
    manual_duplicate = candidate("DOC-2", score=0.90, embedding=[0.99, 0.01])
    ticket = candidate(
        "TICKET-1",
        source=SourceType.TICKET,
        score=0.80,
        embedding=[0, 1],
    )
    result = deduplicate_and_diversify(
        [manual_a, manual_a, manual_duplicate, ticket],
        top_k=3,
    )
    assert len(result) == 2
    assert {item.source_type for item in result} == {SourceType.MANUAL, SourceType.TICKET}


def test_evidence_package_is_stable_bounded_and_excludes_vectors() -> None:
    rewrite = rule_rewrite("PCS 温度高")
    items = [candidate(f"DOC-{index}", embedding=[1.0, 0.0]) for index in range(4)]
    first = build_evidence_package(
        rewrite,
        {"device_type": "PCS"},
        items,
        candidate_counts={"keyword": 4},
        degraded_components=[],
        warnings=[],
    )
    second = build_evidence_package(
        rewrite,
        {"device_type": "PCS"},
        items,
        candidate_counts={"keyword": 4},
        degraded_components=[],
        warnings=[],
    )
    assert first.package_id == second.package_id
    assert len(first.manual_evidence) == 3
    assert "embedding" not in first.model_dump_json()


def test_ticket_embedding_text_and_typed_tool_filters() -> None:
    text = build_ticket_embedding_text(
        {
            "ticket_id": "SHOULD-NOT-APPEAR",
            "device_model": "SC5000",
            "alarm_name": "温度告警",
            "fault_symptom": "温升",
            "root_cause": "",
            "action_taken": "更换风扇",
        }
    )
    assert "SHOULD-NOT-APPEAR" not in text
    context = ToolContext(trace_id="trace", source_system="test")
    manual = ManualSearchInput(
        context=context,
        query="温度",
        filters={"device_type": "PCS", "section_type": ["维护步骤"]},
    )
    ticket = TicketSearchInput(
        context=context,
        query="温度",
        filters={"exclude_ticket_ids": ["T-1"]},
    )
    assert manual.retrieval_mode == "hybrid"
    assert ticket.filters.exclude_ticket_ids == ["T-1"]
    with pytest.raises(ValidationError):
        ManualSearchInput(
            context=context,
            query="温度",
            filters={"unknown": "value"},
        )


def test_phase3_settings_validate_dimension_credentials_and_weight_sums(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    for name in (
        "EMBEDDING_MODE",
        "EMBEDDING_BASE_URL",
        "EMBEDDING_API_KEY",
        "RETRIEVAL_MODE",
    ):
        monkeypatch.delenv(name, raising=False)
    with pytest.raises(ValidationError, match="dimensions"):
        Settings(_env_file=None, embedding_dimension=768)
    with pytest.raises(ValidationError, match="Hybrid retrieval"):
        Settings(_env_file=None, retrieval_mode="hybrid")
    with pytest.raises(ValidationError, match="sum to 1"):
        Settings(_env_file=None, retrieval_keyword_weight=0.5)


def test_openai_strict_schema_requires_nullable_and_defaulted_fields() -> None:
    schema = _openai_strict_schema(QueryRewrite.model_json_schema())
    assert schema["required"] == list(schema["properties"])
    assert schema["additionalProperties"] is False
    assert all("default" not in definition for definition in schema.get("$defs", {}).values())
