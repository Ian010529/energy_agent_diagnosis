from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from energy_agent.core.errors import (
    DocumentParseError,
    DocumentTypeUnsupportedError,
    OcrRequiredError,
)

PARSER_VERSION = "parser.v1.0"


@dataclass(frozen=True)
class ParsedBlock:
    text: str
    page_no: int | None = None
    section_type: str = "正文"


def parse_document(filename: str, content: bytes) -> list[ParsedBlock]:
    suffix = Path(filename).suffix.lower()
    try:
        if suffix in {".txt", ".md", ".markdown"}:
            text = content.decode("utf-8-sig")
            return [ParsedBlock(text=text)] if text.strip() else []
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(content))
            blocks = [
                ParsedBlock(text=text, page_no=index)
                for index, page in enumerate(reader.pages, 1)
                if (text := (page.extract_text() or "").strip())
            ]
            if not blocks:
                raise OcrRequiredError("PDF contains no extractable text; OCR is required")
            return blocks
        if suffix == ".docx":
            document = Document(BytesIO(content))
            paragraphs = []
            for paragraph in document.paragraphs:
                if not paragraph.text.strip():
                    continue
                style_name = paragraph.style.name if paragraph.style else ""
                prefix = "# " if style_name.startswith("Heading") else ""
                paragraphs.append(f"{prefix}{paragraph.text}")
            blocks = [ParsedBlock(text="\n".join(paragraphs))] if paragraphs else []
            for table in document.tables:
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                blocks.append(ParsedBlock(text="\n".join(rows), section_type="表格"))
            return blocks
    except OcrRequiredError:
        raise
    except (UnicodeDecodeError, ValueError, KeyError) as exc:
        raise DocumentParseError(f"Unable to parse {suffix} document") from exc
    raise DocumentTypeUnsupportedError(f"Unsupported document type: {suffix}")
